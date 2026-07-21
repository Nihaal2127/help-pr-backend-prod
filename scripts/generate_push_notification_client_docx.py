"""Generate client-shareable push notification specification (.docx)."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent.parent / "docs" / "Mobile_Push_Notifications_Client_Spec.docx"

HEADER_FILL = "2F5496"
ALT_ROW_FILL = "D6E4F0"


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def style_header_row(table, fill_hex: str = HEADER_FILL) -> None:
    for cell in table.rows[0].cells:
        set_cell_shading(cell, fill_hex)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
                run.font.name = "Calibri"


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
    style_header_row(table)

    for row_idx, row_data in enumerate(rows):
        row = table.add_row().cells
        for i, text in enumerate(row_data):
            row[i].text = str(text)
            for paragraph in row[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
        if row_idx % 2 == 1:
            for cell in row:
                set_cell_shading(cell, ALT_ROW_FILL)

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)
    return table


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(47, 84, 150)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


PLACEHOLDERS = [
    ("{order_id}", "ORD-1024", "Order reference number"),
    ("{quote_id}", "QT-558", "Quote reference number"),
    ("{service_name}", "AC Repair", "Name of the service on the order"),
    ("{amount}", "1500.00", "Monetary amount (2 decimal places)"),
    ("{charge_label}", "Extra parts", "Label for an additional service charge"),
    ("{total_amount}", "1770.00", "Charge total including tax"),
    ("{new_due_amount}", "2500.00", "Updated amount the customer needs to pay"),
    ("{payer_type}", "customer / partner", "Who made the payment (when applicable)"),
    ("{status}", "accepted, In-Progress, completed", "New status value"),
    ("{plan_name}", "Gold Plan", "Subscription plan name"),
    ("{description}", "Order payment credit", "Short wallet transaction description"),
    ("{dispute_id}", "DSP-12", "Dispute reference number"),
    ("{ticket_id}", "TKT-88", "Support ticket reference number"),
    ("{sender_name}", "John", "Name of the person who sent a chat message"),
    ("{message_preview}", "Hello, when will you arrive?", "Chat message text or attachment summary"),
]

MASTER_ROWS = [
    # Quotes
    ("Q1", "Quotes", "New quote created", "Yes", "Yes (if assigned)", "New quote", "Quote #{quote_id} has been created.", "Available"),
    ("Q2", "Quotes", "Customer or admin assigns a partner to a quote", "No", "Yes (assigned partner)", "New quote request", "Quote #{quote_id} has been assigned to you. Please review.", "Available"),
    ("Q3", "Quotes", "Partner accepts quote", "Yes", "Yes", "Quote status update", "Quote #{quote_id} status changed to accepted.", "Available"),
    ("Q4", "Quotes", "Partner rejects quote", "Yes", "Yes", "Quote status update", "Quote #{quote_id} status changed to failed.", "Available"),
    ("Q5", "Quotes", "Customer cancels quote", "No", "Yes", "Quote status update", "Quote #{quote_id} status changed to failed.", "Available"),
    ("Q6", "Quotes", "Quote status updated by admin", "Yes", "Yes", "Quote status update", "Quote #{quote_id} status changed to {status}.", "Available"),
    ("Q7", "Quotes", "Quote converted to order", "Yes", "Yes", "Quote status update", "Quote #{quote_id} status changed to success.", "Available"),
    ("Q8", "Quotes", "Customer converts quote to order (mobile)", "Yes", "Yes", "Quote status update", "Quote #{quote_id} status changed to success.", "Available"),
    # Orders
    ("O1", "Orders", "New order created", "Yes", "Yes", "New order", "Order #{order_id} has been created.", "Available"),
    ("O2", "Orders", "Order status changed", "Yes", "Yes", "Order status update", "Order #{order_id} status changed to {status}.", "Available"),
    ("O3", "Orders", "Order cancelled", "Yes", "Yes", "Order cancelled", "Order #{order_id} has been cancelled.", "Available"),
    ("O4", "Orders", "Order refunded (status → Refunded)", "Yes", "Yes", "Order status update", "Order #{order_id} status changed to Refunded.", "Coming soon"),
    ("O5", "Orders", "Service line status changed", "Yes", "Yes", "Service update", "{service_name} status changed to {status} for order #{order_id}.", "Available"),
    ("O6", "Orders", "Partner assigned to a service", "No", "Yes (assigned partner)", "New service assigned", "You have a new service ({service_name}) for order #{order_id}.", "Available"),
    ("O7", "Orders", "Partner removed from a service", "No", "Yes (previous partner)", "Service cancelled", "Service for order #{order_id} has been removed from your list.", "Available"),
    ("O8", "Orders", "Service date or time updated", "No", "Yes (assigned partner)", "Service time updated", "Time updated for service ({service_name}) of order #{order_id}.", "Available"),
    ("O9", "Orders", "Single service line cancelled", "Yes", "Yes", "Service cancelled", "Your {service_name} for order #{order_id} has been cancelled.", "Available"),
    ("O11", "Orders", "Partner starts work on order", "Yes", "No", "Partner on the way", "Your partner has started work on order #{order_id}.", "Available"),
    ("O12", "Orders", "Partner completes order", "Yes", "No", "Order completed", "Order #{order_id} has been completed by your partner.", "Available"),
    # Additional charges
    ("AC1", "Additional service charges", "Partner adds an additional charge", "Yes", "Yes", "Additional charge added", "Additional charge of {amount}{charge_label} added to order #{order_id}.", "Available"),
    ("AC2", "Additional service charges", "Admin adds an additional charge", "Yes", "Yes", "Additional charge added", "Additional charge of {amount}{charge_label} added to order #{order_id}.", "Available"),
    ("AC3", "Additional service charges", "Additional charge added when order is created", "Yes", "Yes", "Additional charge added", "Additional charge of {amount}{charge_label} added to order #{order_id}.", "Available"),
    ("AC4", "Additional service charges", "Additional charge updated", "Yes", "Yes", "Additional charge updated", "Additional charge {charge_label} on order #{order_id} was updated to {total_amount}.", "Available"),
    ("AC5", "Additional service charges", "Additional charge removed", "Yes", "Yes", "Additional charge removed", "Additional charge {charge_label} was removed from order #{order_id}.", "Available"),
    ("AC6", "Additional service charges", "Customer reminded to pay updated balance", "Yes", "No", "Payment due updated", "An additional charge of {total_amount}{charge_label} was added to order #{order_id}. Amount due: {new_due_amount}.", "Coming soon"),
    # Payments
    ("P1", "Payments", "Payment received on order", "Yes", "Yes", "Payment received", "Payment of {amount} received for order #{order_id}{payer_type}.", "Available"),
    ("P2", "Payments", "Payment on quote to order conversion", "Yes", "Yes", "Payment received", "Payment of {amount} received for order #{order_id}.", "Available"),
    ("P3", "Payments", "Online payment failed", "Yes", "No", "Payment failed", "Payment for order #{order_id} could not be completed. Please try again.", "Available"),
    ("P4", "Payments", "Refund processed", "Yes", "Yes", "Refund processed", "A refund of {amount} has been processed for order #{order_id}.", "Available"),
    # Wallet
    ("W1", "Wallet", "Wallet credited", "No", "Yes", "Wallet credit", "{amount} credited to your wallet{description}.", "Available"),
    ("W2", "Wallet", "Wallet debited", "No", "Yes", "Wallet debit", "{amount} debited from your wallet{description}.", "Available"),
    ("W3", "Wallet", "Wallet debited due to refund", "No", "Yes", "Wallet debit", "{amount} debited from your wallet{description}.", "Available"),
    # Subscriptions
    ("S1", "Subscriptions", "Subscription plan assigned", "No", "Yes", "Subscription assigned", 'Subscription plan "{plan_name}" has been assigned to you.', "Available"),
    ("S2", "Subscriptions", "Subscription status changed", "No", "Yes", "Subscription update", "Your subscription status is now {status}.", "Available"),
    ("S3", "Subscriptions", "Partner changes plan (upgrade/downgrade)", "No", "Yes", "Subscription update", 'Your subscription plan has been changed to "{plan_name}".', "Available"),
    ("S4", "Subscriptions", "Subscription payment successful", "No", "Yes", "Subscription update", 'Your subscription payment was successful. Plan: "{plan_name}".', "Available"),
    # Disputes
    ("D2", "Disputes", "Dispute status updated", "Yes", "No", "Dispute update", "Your dispute {dispute_id} for order #{order_id} is now {status}.", "Available"),
    # Chat
    ("C1", "Chat", "New chat message (recipient offline)", "Yes", "Yes", "{sender_name}", "{message_preview}", "Available"),
    # Tickets
    ("T1", "Support", "Support ticket status changed", "Yes", "No", "Ticket Update", "Your ticket {ticket_id} status changed to {status}.", "Available"),
    # Account
    ("A1", "Account", "Partner account verified", "No", "Yes", "Account verified", "Your partner account has been verified. You can now accept jobs.", "Available"),
    ("A2", "Account", "Partner verification not approved", "No", "Yes", "Verification update", "Your partner verification was not approved. Please check your documents.", "Available"),
    # Appointments
    ("AP1", "Appointments", "Appointment scheduled (manual or auto on order create)", "Yes", "Yes", "Appointment scheduled", "A service appointment has been scheduled for order #{order_id}.", "Available"),
    ("AP2", "Appointments", "Appointment status changed", "Yes", "Yes", "Appointment update", "Your appointment for order #{order_id} is now {status}.", "Available"),
    # Reviews
    ("R1", "Reviews", "Customer submits order review", "No", "Yes", "New review", "You received a new review for order #{order_id}.", "Available"),
    # Reminders
    ("RM1", "Reminders", "Upcoming service reminder", "Yes", "Yes", "Service reminder", "Your service for order #{order_id} is scheduled soon.", "Available"),
    ("RM2", "Reminders", "Quote pending action reminder", "Yes", "Yes", "Action required", "Quote #{quote_id} is waiting for your response.", "Available"),
    ("RM3", "Reminders", "Subscription expiring soon", "No", "Yes", "Subscription reminder", "Your subscription plan expires soon. Renew to continue receiving jobs.", "Available"),
]

CUSTOMER_ROWS = [
    ("Q1", "New quote created", "New quote", "Quote #{quote_id} has been created."),
    ("Q3", "Partner accepted quote", "Quote status update", "Quote #{quote_id} status changed to accepted."),
    ("Q4", "Partner rejected quote", "Quote status update", "Quote #{quote_id} status changed to failed."),
    ("Q6", "Quote status updated", "Quote status update", "Quote #{quote_id} status changed to {status}."),
    ("Q7", "Quote converted to order", "Quote status update", "Quote #{quote_id} status changed to success."),
    ("Q8", "Customer converts quote to order", "Quote status update", "Quote #{quote_id} status changed to success."),
    ("O1", "New order created", "New order", "Order #{order_id} has been created."),
    ("O2", "Order status changed", "Order status update", "Order #{order_id} status changed to {status}."),
    ("O3", "Order cancelled", "Order cancelled", "Order #{order_id} has been cancelled."),
    ("O4", "Order refunded", "Order status update", "Order #{order_id} status changed to Refunded."),
    ("O5", "Service status changed", "Service update", "{service_name} status changed to {status} for order #{order_id}."),
    ("O9", "Service line cancelled", "Service cancelled", "Your {service_name} for order #{order_id} has been cancelled."),
    ("AC1", "Partner added additional charge", "Additional charge added", "Additional charge of {amount}{charge_label} added to order #{order_id}."),
    ("AC2", "Admin added additional charge", "Additional charge added", "Additional charge of {amount}{charge_label} added to order #{order_id}."),
    ("AC3", "Additional charge on order create", "Additional charge added", "Additional charge of {amount}{charge_label} added to order #{order_id}."),
    ("AC4", "Additional charge updated", "Additional charge updated", "Additional charge {charge_label} on order #{order_id} was updated to {total_amount}."),
    ("AC5", "Additional charge removed", "Additional charge removed", "Additional charge {charge_label} was removed from order #{order_id}."),
    ("AC6", "Payment due after new charge", "Payment due updated", "An additional charge of {total_amount}{charge_label} was added to order #{order_id}. Amount due: {new_due_amount}."),
    ("O11", "Partner started work", "Partner on the way", "Your partner has started work on order #{order_id}."),
    ("O12", "Partner completed order", "Order completed", "Order #{order_id} has been completed by your partner."),
    ("P1", "Payment received", "Payment received", "Payment of {amount} received for order #{order_id}{payer_type}."),
    ("P2", "Payment on quote conversion", "Payment received", "Payment of {amount} received for order #{order_id}."),
    ("P3", "Payment failed", "Payment failed", "Payment for order #{order_id} could not be completed. Please try again."),
    ("P4", "Refund processed", "Refund processed", "A refund of {amount} has been processed for order #{order_id}."),
    ("D2", "Dispute status updated", "Dispute update", "Your dispute {dispute_id} for order #{order_id} is now {status}."),
    ("C1", "New chat message", "{sender_name}", "{message_preview}"),
    ("T1", "Support ticket updated", "Ticket Update", "Your ticket {ticket_id} status changed to {status}."),
    ("AP1", "Appointment scheduled (manual or auto)", "Appointment scheduled", "A service appointment has been scheduled for order #{order_id}."),
    ("AP2", "Appointment updated", "Appointment update", "Your appointment for order #{order_id} is now {status}."),
    ("RM1", "Service reminder", "Service reminder", "Your service for order #{order_id} is scheduled soon."),
    ("RM2", "Quote action reminder", "Action required", "Quote #{quote_id} is waiting for your response."),
]

PARTNER_ROWS = [
    ("Q1", "New quote created", "New quote", "Quote #{quote_id} has been created."),
    ("Q2", "Quote assigned by customer or admin", "New quote request", "Quote #{quote_id} has been assigned to you. Please review."),
    ("Q3", "Quote accepted", "Quote status update", "Quote #{quote_id} status changed to accepted."),
    ("Q4", "Quote rejected", "Quote status update", "Quote #{quote_id} status changed to failed."),
    ("Q5", "Customer cancelled quote", "Quote status update", "Quote #{quote_id} status changed to failed."),
    ("Q6", "Quote status updated", "Quote status update", "Quote #{quote_id} status changed to {status}."),
    ("Q7", "Quote converted to order", "Quote status update", "Quote #{quote_id} status changed to success."),
    ("Q8", "Customer converted quote to order", "Quote status update", "Quote #{quote_id} status changed to success."),
    ("O1", "New order created", "New order", "Order #{order_id} has been created."),
    ("O2", "Order status changed", "Order status update", "Order #{order_id} status changed to {status}."),
    ("O3", "Order cancelled", "Order cancelled", "Order #{order_id} has been cancelled."),
    ("O4", "Order refunded", "Order status update", "Order #{order_id} status changed to Refunded."),
    ("O5", "Service status changed", "Service update", "{service_name} status changed to {status} for order #{order_id}."),
    ("O6", "New service assigned", "New service assigned", "You have a new service ({service_name}) for order #{order_id}."),
    ("O7", "Service removed from partner", "Service cancelled", "Service for order #{order_id} has been removed from your list."),
    ("O8", "Service time updated", "Service time updated", "Time updated for service ({service_name}) of order #{order_id}."),
    ("O9", "Service line cancelled", "Service cancelled", "Your {service_name} for order #{order_id} has been cancelled."),
    ("AC1", "Partner added additional charge", "Additional charge added", "Additional charge of {amount}{charge_label} added to order #{order_id}."),
    ("AC2", "Admin added additional charge", "Additional charge added", "Additional charge of {amount}{charge_label} added to order #{order_id}."),
    ("AC3", "Additional charge on order create", "Additional charge added", "Additional charge of {amount}{charge_label} added to order #{order_id}."),
    ("AC4", "Additional charge updated", "Additional charge updated", "Additional charge {charge_label} on order #{order_id} was updated to {total_amount}."),
    ("AC5", "Additional charge removed", "Additional charge removed", "Additional charge {charge_label} was removed from order #{order_id}."),
    ("P1", "Payment received on order", "Payment received", "Payment of {amount} received for order #{order_id}{payer_type}."),
    ("P2", "Payment on quote conversion", "Payment received", "Payment of {amount} received for order #{order_id}."),
    ("P4", "Refund processed", "Refund processed", "A refund of {amount} has been processed for order #{order_id}."),
    ("W1", "Wallet credited", "Wallet credit", "{amount} credited to your wallet{description}."),
    ("W2", "Wallet debited", "Wallet debit", "{amount} debited from your wallet{description}."),
    ("W3", "Wallet debited (refund)", "Wallet debit", "{amount} debited from your wallet: Refund adjustment for order #{order_id}."),
    ("S1", "Subscription assigned", "Subscription assigned", 'Subscription plan "{plan_name}" has been assigned to you.'),
    ("S2", "Subscription status changed", "Subscription update", "Your subscription status is now {status}."),
    ("S3", "Plan changed", "Subscription update", 'Your subscription plan has been changed to "{plan_name}".'),
    ("S4", "Subscription payment success", "Subscription update", 'Your subscription payment was successful. Plan: "{plan_name}".'),
    ("C1", "New chat message", "{sender_name}", "{message_preview}"),
    ("A1", "Account verified", "Account verified", "Your partner account has been verified. You can now accept jobs."),
    ("A2", "Verification not approved", "Verification update", "Your partner verification was not approved. Please check your documents."),
    ("AP1", "Appointment scheduled (manual or auto)", "Appointment scheduled", "A service appointment has been scheduled for order #{order_id}."),
    ("AP2", "Appointment updated", "Appointment update", "Your appointment for order #{order_id} is now {status}."),
    ("R1", "New customer review", "New review", "You received a new review for order #{order_id}."),
    ("RM1", "Service reminder", "Service reminder", "Your service for order #{order_id} is scheduled soon."),
    ("RM3", "Subscription expiring", "Subscription reminder", "Your subscription plan expires soon. Renew to continue receiving jobs."),
]

CHARGE_EXAMPLES = [
    ("Partner adds charge: Extra pipe fitting (500.00)", "Additional charge added", "Additional charge of 500.00 (Extra pipe fitting) added to order #ORD-1024."),
    ("Admin adds charge with no label (200.00)", "Additional charge added", "Additional charge of 200.00 added to order #ORD-1024."),
    ("Charge added when order is created (150.00 Transport)", "Additional charge added", "Additional charge of 150.00 (Transport) added to order #ORD-1024."),
    ("Customer payment due after new charge", "Payment due updated", "An additional charge of 590.00 (Labour) was added to order #ORD-1024. Amount due: 2500.00."),
]


def build_document() -> Document:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Mobile Push Notifications\nSpecification")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(47, 84, 150)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Customer App & Partner App")
    r2.font.size = Pt(14)
    r2.font.name = "Calibri"
    r2.font.color.rgb = RGBColor(89, 89, 89)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = date_p.add_run("July 2026")
    r3.font.size = Pt(11)
    r3.font.name = "Calibri"
    r3.italic = True

    doc.add_paragraph()

    add_heading(doc, "1. Overview", 1)
    add_body(
        doc,
        "This document defines all push notification instances for the Help platform mobile applications. "
        "For each event it specifies when the notification is sent, which app receives it (Customer or Partner), "
        "and the notification title and message template.",
    )
    add_body(
        doc,
        "Placeholders in curly braces (e.g. {order_id}) are replaced with real values when the notification is sent.",
    )

    add_heading(doc, "2. Applications", 1)
    add_table(
        doc,
        ["Application", "Users", "Receives push notifications"],
        [
            ("Customer mobile app", "End customers who book services", "Yes"),
            ("Partner mobile app", "Service partners who fulfil orders", "Yes"),
        ],
        col_widths=[2.0, 2.5, 1.8],
    )
    doc.add_paragraph()
    add_body(
        doc,
        "General rules: The user who performed the action does not receive the same notification. "
        "Users can turn off update notifications and reminder notifications separately in app settings.",
    )

    add_heading(doc, "3. Template placeholders", 1)
    add_table(
        doc,
        ["Placeholder", "Example", "Description"],
        PLACEHOLDERS,
        col_widths=[1.3, 1.5, 3.5],
    )

    doc.add_page_break()

    add_heading(doc, "4. Complete notification matrix", 1)
    add_body(
        doc,
        "Master list of all notification instances. Status: Available = live in the app today; Coming soon = planned for a future release.",
    )

    # Landscape for wide table
    new_section = doc.add_section()
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    new_section.left_margin = Inches(0.5)
    new_section.right_margin = Inches(0.5)

    add_table(
        doc,
        [
            "Ref",
            "Module",
            "When notification is sent",
            "Customer app",
            "Partner app",
            "Title",
            "Message template",
            "Status",
        ],
        MASTER_ROWS,
        col_widths=[0.45, 1.1, 1.8, 0.7, 0.85, 1.0, 2.5, 0.75],
    )

    doc.add_page_break()

    # Portrait again
    portrait = doc.add_section()
    portrait.orientation = WD_ORIENT.PORTRAIT
    portrait.page_width, portrait.page_height = portrait.page_height, portrait.page_width

    add_heading(doc, "5. Customer mobile app notifications", 1)
    add_body(doc, "All notifications a customer may receive, grouped by reference ID from Section 4.")
    add_table(
        doc,
        ["Ref", "When notification is sent", "Title", "Message template"],
        CUSTOMER_ROWS,
        col_widths=[0.5, 2.0, 1.3, 3.5],
    )

    doc.add_page_break()

    add_heading(doc, "6. Partner mobile app notifications", 1)
    add_body(doc, "All notifications a partner may receive, grouped by reference ID from Section 4.")
    add_table(
        doc,
        ["Ref", "When notification is sent", "Title", "Message template"],
        PARTNER_ROWS,
        col_widths=[0.5, 2.0, 1.3, 3.5],
    )

    doc.add_page_break()

    add_heading(doc, "7. Additional service charges", 1)
    add_body(
        doc,
        "Additional service charges are extra fees added to an order (e.g. materials, labour, transport). "
        "They increase the order total and may change the amount the customer needs to pay.",
    )
    add_heading(doc, "7.1 Notification instances", 2)
    charge_rows = [r for r in MASTER_ROWS if r[1] == "Additional service charges"]
    add_table(
        doc,
        ["Ref", "When notification is sent", "Customer app", "Partner app", "Title", "Message template", "Status"],
        [(r[0], r[2], r[3], r[4], r[5], r[6], r[7]) for r in charge_rows],
        col_widths=[0.5, 2.0, 0.8, 0.8, 1.2, 2.8, 0.9],
    )
    doc.add_paragraph()

    add_heading(doc, "7.2 Example messages", 2)
    add_table(
        doc,
        ["Scenario", "Title", "Message"],
        CHARGE_EXAMPLES,
        col_widths=[2.2, 1.3, 3.8],
    )

    doc.add_page_break()

    add_heading(doc, "8. Chat notifications", 1)
    add_body(
        doc,
        "When a user receives a new chat message while offline, a push notification is sent with the sender's name as the title "
        "and the message text (or 'Sent an image' / 'Sent a file' for attachments) as the body. "
        "Tapping the notification opens the relevant chat thread.",
    )

    add_heading(doc, "9. Status summary", 1)
    available = sum(1 for r in MASTER_ROWS if r[7] == "Available")
    coming = sum(1 for r in MASTER_ROWS if r[7] == "Coming soon")
    add_table(
        doc,
        ["Category", "Count"],
        [
            ("Total notification instances", str(len(MASTER_ROWS))),
            ("Available now", str(available)),
            ("Coming soon", str(coming)),
            ("Customer app instances (Section 5)", str(len(CUSTOMER_ROWS))),
            ("Partner app instances (Section 6)", str(len(PARTNER_ROWS))),
        ],
        col_widths=[3.5, 1.5],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph("— End of document —")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

    return doc


def main():
    doc = build_document()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
